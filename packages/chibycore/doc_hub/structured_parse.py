"""结构化解析：Markdown / 纯文本 / docx 标题树；PDF 字号启发式。"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple

from chibycore.doc_hub.chunker_v2 import ParsedDocument, Section
from chibycore.doc_hub.parse import SUPPORTED_SUFFIXES, parse_file, sniff_title

_MD_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def parse_markdown_sections(text: str, *, doc_title: str = "") -> ParsedDocument:
    lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    root_children: List[Section] = []
    stack: List[Section] = []  # 当前路径
    preamble: List[str] = []

    def ensure_level(level: int) -> None:
        while stack and stack[-1].level >= level:
            stack.pop()

    for line in lines:
        m = _MD_HEADING.match(line)
        if m:
            # flush preamble into synthetic section
            if preamble and not stack and not root_children:
                body = "\n".join(preamble).strip()
                if body:
                    root_children.append(
                        Section(title=doc_title or "正文", level=1, content=body)
                    )
                preamble = []
            level = len(m.group(1))
            title = m.group(2).strip()
            sec = Section(title=title, level=level, content="")
            ensure_level(level)
            if stack:
                stack[-1].children.append(sec)
            else:
                root_children.append(sec)
            stack.append(sec)
        else:
            if stack:
                stack[-1].content += line + "\n"
            else:
                preamble.append(line)

    if preamble and not root_children:
        body = "\n".join(preamble).strip()
        root_children.append(
            Section(title=doc_title or "正文", level=1, content=body or text)
        )
    elif preamble and root_children and not any(s.content.strip() for s in root_children[:1]):
        # 标题前导文挂到第一个节
        lead = "\n".join(preamble).strip()
        if lead:
            root_children[0].content = lead + "\n" + root_children[0].content

    if not root_children:
        root_children = [Section(title=doc_title or "正文", level=1, content=text or "")]

    return ParsedDocument(title=doc_title or sniff_title(Path("x.md"), text), sections=root_children)


def parse_docx_sections(path: Path, *, doc_title: str = "") -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    root_children: List[Section] = []
    stack: List[Section] = []
    buf: List[str] = []

    def flush_buf_to_current() -> None:
        if not buf:
            return
        text = "\n".join(buf).strip()
        buf.clear()
        if not text:
            return
        if stack:
            stack[-1].content += text + "\n"
        else:
            root_children.append(Section(title=doc_title or "正文", level=1, content=text))

    def heading_level(style_name: str) -> int:
        name = (style_name or "").lower()
        for i in range(1, 7):
            if f"heading {i}" in name or f"标题 {i}" in name or name == f"heading{i}":
                return i
        return 0

    for para in doc.paragraphs:
        text = (para.text or "").rstrip()
        style = ""
        try:
            style = para.style.name if para.style else ""
        except Exception:
            style = ""
        lvl = heading_level(style)
        if lvl > 0 and text.strip():
            flush_buf_to_current()
            sec = Section(title=text.strip(), level=lvl, content="")
            while stack and stack[-1].level >= lvl:
                stack.pop()
            if stack:
                stack[-1].children.append(sec)
            else:
                root_children.append(sec)
            stack.append(sec)
        else:
            if text.strip():
                buf.append(text)
    flush_buf_to_current()
    if not root_children:
        flat = "\n".join(p.text for p in doc.paragraphs)
        root_children = [Section(title=doc_title or "正文", level=1, content=flat)]
    return ParsedDocument(title=doc_title or path.stem, sections=root_children)


def parse_to_document(path: Path) -> Tuple[ParsedDocument, str]:
    """解析为结构化文档；同时返回原始纯文本（用于长度门禁）。"""
    p = Path(path)
    plain, title = parse_file(p)
    suf = p.suffix.lower()
    if suf in (".md", ".markdown"):
        doc = parse_markdown_sections(plain, doc_title=title)
        if not doc.structure_quality:
            n = len(doc.sections or [])
            has_child = any(s.children for s in (doc.sections or []))
            doc.structure_quality = (
                "high" if has_child or n >= 3 else ("medium" if n >= 2 else "low")
            )
        return doc, plain
    if suf == ".docx":
        try:
            doc = parse_docx_sections(p, doc_title=title)
            if not doc.structure_quality:
                doc.structure_quality = (
                    "high" if any(s.children for s in doc.sections) else "medium"
                )
            return doc, plain
        except Exception:
            return (
                ParsedDocument(
                    title=title,
                    sections=[Section(title=title or "正文", level=1, content=plain)],
                    structure_quality="low",
                ),
                plain,
            )
    if suf == ".pdf":
        try:
            from chibycore.doc_hub.pdf_structure import parse_pdf_sections

            doc = parse_pdf_sections(p, doc_title=title)
            return doc, plain
        except Exception:
            return (
                ParsedDocument(
                    title=title,
                    sections=[Section(title=title or "正文", level=1, content=plain)],
                    structure_quality="low",
                ),
                plain,
            )
    # txt：扁平
    return (
        ParsedDocument(
            title=title,
            sections=[Section(title=title or "正文", level=1, content=plain)],
            structure_quality="low",
        ),
        plain,
    )


# re-export for callers that only need suffixes
__all__ = [
    "SUPPORTED_SUFFIXES",
    "parse_to_document",
    "parse_markdown_sections",
    "parse_docx_sections",
    "ParsedDocument",
    "Section",
]
